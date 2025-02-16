import docx
import mammoth
import pdfkit

class MarkdownOutput:
    def generate_markdown_output(self, template, minutes_content, file_name):
        agenda = minutes_content["agenda"]
        participants = self.generate_list(minutes_content["participants"])
        summary = minutes_content["summary"]
        action_items = self.generate_table(minutes_content["action_items"])
        meeting_minutes = self.generate_table(minutes_content["meeting_minutes"])

        output = template.format(agenda=agenda, participants=participants, summary=summary, action_items=action_items, meeting_minutes=meeting_minutes)
        with open(file_name, "w") as file:
            file.write(output)

        print("Meeting minutes saved to meeting_minutes.md")

    def generate_table(self, obj):
        if not obj:
            return ""
        
        keys = list(obj[0].keys())
        header = f"| {' | '.join(keys)} |"
        table_content = [[] for _ in range(len(obj))]
        
        for key in keys:
            for i in range(len(obj)):
                table_content[i].append(obj[i][key])
        
        for i in range(len(obj)):
            table_content[i] = f"| {' | '.join(table_content[i])} |"
        
        return "\n".join([header] + [f"| {' | '.join(['-' for i in range(len(keys))])} |"] + table_content)
    
    def generate_list(self, arr):
        if not arr:
            return ""
        
        return "\n".join([f"+ {item}" for item in arr])

class DOCXOutput:
    def __init__(self):
        self.doc = None
    
    def generate_docx_output(self, template, minutes_content, file_name):
        self.doc = docx.Document()
        tokens = template.split("\n")

        for token in tokens:
            if token.startswith("#"):
                self.doc.add_heading(token[2:], level=1)
            elif token.startswith("##"):
                self.doc.add_heading(token[3:], level=1)
            elif token.startswith("{"):
                key = token[1:-1]
                if key == "action_items" and minutes_content["action_items"]:
                    header = {"action_item": "Action Item", "assignee": "Assignee", "due_date": "Due Date"}
                    self.generate_table(minutes_content["action_items"], header)
                elif key == "meeting_minutes" and minutes_content["meeting_minutes"]:
                    header = {"minutes": "Minutes", "speaker": "Speaker", "content": "Content"}
                    self.generate_table(minutes_content["meeting_minutes"], header)
                elif key == "participants" and minutes_content["participants"]:
                    self.generate_list(minutes_content["participants"])
                else:
                    self.doc.add_paragraph(minutes_content[key], style="BodyText")
            else:
                self.doc.add_paragraph(token, style="BodyText")
        
        self.doc.save(file_name)

    def generate_table(self, obj, header):
        table = self.doc.add_table(rows=1, cols=len(obj[0]))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, key in enumerate(header.keys()):
            hdr_cells[i].text = header[key]
        
        for item in obj:
            row_cells = table.add_row().cells
            for i, key in enumerate(obj[0].keys()):
                row_cells[i].text = item[key]
    
    def generate_list(self, arr):
        for i in range(len(arr)):
            self.doc.add_paragraph(arr[i], style="ListBullet")
        
    
class PDFOutput(DOCXOutput):
    def generate_pdf_output(self, template, minutes_content, file_name, docx_file_name):
        self.generate_docx_output(template, minutes_content, docx_file_name)
        with open(docx_file_name, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            pdfkit.from_string(html, file_name)

if __name__ == "__main__":
    markdown_output = MarkdownOutput()
    docx_output = DOCXOutput()
    pdf_output = PDFOutput()

    with open("test_transcript.txt", "r") as file:
        import json
        from models import MeetingOfMinutesLLM

        llm = MeetingOfMinutesLLM()
        transcript = file.read()
        minutes_content = json.loads(llm.generate_minutes(transcript))
        
        with open("./templates/sample_template.md", "r") as file:
            template = file.read()
            markdown_output.generate_markdown_output(template, minutes_content, "meeting_minutes.md")
            docx_output.generate_docx_output(template, minutes_content, "meeting_minutes.docx")
            pdf_output.generate_pdf_output(template, minutes_content, "meeting_minutes.pdf", "meeting_minutes.docx")